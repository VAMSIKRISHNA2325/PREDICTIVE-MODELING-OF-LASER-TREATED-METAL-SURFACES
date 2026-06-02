"""
generate_report.py — Build the detailed project report (.docx)

Pulls live metrics from the trained models and embeds figures + real code
snapshots so the report always matches the current code.

Run:  python generate_report.py    →    Project_Report.docx
"""

import warnings, os, sys, re
warnings.filterwarnings("ignore")
os.environ["PYTHONWARNINGS"] = "ignore"

import numpy as np
import pandas as pd
import joblib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sklearn.model_selection import KFold, LeaveOneOut, cross_val_predict
from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error

from config import (DATA_DIR, MODEL_DIR, BASE_DIR, FEATURES, BASE_FEATURES,
                    DERIVED_FEATURES, FEATURE_LABELS, PARAM_RANGES, TARGETS,
                    SEED, N_ESTIMATORS, RF_PARAMS, EXTRATREES_PARAMS,
                    add_derived_features, build_model)
import report_figures

for _s in (sys.stdout, sys.stderr):
    try:
        _s.reconfigure(encoding="utf-8")
    except (AttributeError, ValueError):
        pass

RESULTS = BASE_DIR / "results"

# ── Generate figures + before/after numbers ──────────────────────────────────
BA = report_figures.make_all()  # {key: {label, unit, before:(r2,mae), after:(r2,mae)}}


# ── Live evaluation of the deployed models ───────────────────────────────────
def evaluate(key, spec):
    df = add_derived_features(pd.read_csv(DATA_DIR / spec["csv"]))
    X, y = df[FEATURES], df[spec["column"]].values
    cv = LeaveOneOut() if len(df) < 40 else KFold(5, shuffle=True, random_state=SEED)
    cvname = "Leave-One-Out" if len(df) < 40 else "5-fold"
    yp = cross_val_predict(build_model(spec["model"]), X, y, cv=cv, n_jobs=-1)
    model = joblib.load(MODEL_DIR / f"model_{key}.pkl")
    return {
        "rows": len(df), "cv": cvname,
        "r2": r2_score(y, yp), "mae": mean_absolute_error(y, yp),
        "rmse": np.sqrt(mean_squared_error(y, yp)),
        "range": (float(y.min()), float(y.max())),
        "importance": dict(zip(FEATURES, model.named_steps["model"].feature_importances_)),
        "model": model,
    }


EV = {k: evaluate(k, v) for k, v in TARGETS.items()}

# ── Document setup ───────────────────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)


def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def p(t): return doc.add_paragraph(t)
def bullet(t): doc.add_paragraph(t, style="List Bullet")


def caption(t):
    par = doc.add_paragraph(t); par.runs[0].italic = True; return par


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for c, htext in zip(t.rows[0].cells, headers):
        c.text = str(htext); c.paragraphs[0].runs[0].font.bold = True
    for r in rows:
        for c, v in zip(t.add_row().cells, r):
            c.text = str(v)
    return t


def code_block(text):
    """Render source code in a monospace, lightly shaded paragraph."""
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.font.name = "Consolas"; run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    par.paragraph_format.left_indent = Inches(0.2)
    par.paragraph_format.space_before = Pt(2); par.paragraph_format.space_after = Pt(6)
    return par


def snippet(filename, start_marker, end_marker=None, max_lines=40):
    """Extract a code snippet from a project file between markers (inclusive
    start, exclusive end). Falls back to the first max_lines."""
    src = (BASE_DIR / filename).read_text(encoding="utf-8").splitlines()
    i0 = next((i for i, l in enumerate(src) if start_marker in l), 0)
    if end_marker:
        i1 = next((i for i in range(i0 + 1, len(src)) if end_marker in src[i]), len(src))
    else:
        i1 = min(i0 + max_lines, len(src))
    return "\n".join(src[i0:i1])


def add_picture(path, width=6.0):
    if Path_exists(path):
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def Path_exists(path):
    from pathlib import Path
    return Path(path).exists()


# ── Title page ───────────────────────────────────────────────────────────────
t = doc.add_heading("Prediction of Surface Roughness and Hardness Using "
                    "Machine Learning", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
p("A Detailed Machine-Learning Project Report").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Abstract ─────────────────────────────────────────────────────────────────
h1("Abstract")
p("This project develops machine-learning models that predict two key part "
  "properties — surface roughness and hardness — directly from four laser "
  "process parameters: laser power, scan speed, hatch distance, and layer "
  "thickness. Experimental data were compiled from the published literature "
  "[1]–[16] and used to train two independent regression models. Two engineered "
  "features (an energy term and a power-to-speed ratio) were added, and the "
  "best-performing algorithm was selected separately for each property. Through "
  "feature engineering and model selection, the cross-validated accuracy of the "
  f"hardness model improved markedly from R² = {BA['hardness']['before'][0]:.2f} "
  f"to R² = {BA['hardness']['after'][0]:.2f}, while surface roughness improved "
  f"from R² = {BA['roughness']['before'][0]:.2f} to "
  f"R² = {BA['roughness']['after'][0]:.2f}. The trained models are deployed in an "
  "interactive application that returns predictions with an uncertainty estimate.")
pr = p("Keywords: surface roughness, hardness, Random Forest, Extra Trees, feature "
       "engineering, regression, machine learning."); pr.runs[0].italic = True

# ── List of Symbols and Abbreviations ────────────────────────────────────────
h1("List of Symbols and Abbreviations")
caption("Symbols and abbreviations used in this report.")
table(["Symbol / Abbreviation", "Meaning", "Unit"], [
    ["P", "Laser power", "W"],
    ["v", "Scan speed", "mm/s"],
    ["h", "Hatch distance", "mm"],
    ["t", "Layer thickness", "mm"],
    ["E", "Energy term, P / (v · h · t)", "—"],
    ["P/v", "Power-to-speed ratio (derived feature)", "—"],
    ["SR", "Surface roughness", "µm"],
    ["HV", "Hardness (Vickers)", "HV"],
    ["ML", "Machine learning", "—"],
    ["RF", "Random Forest", "—"],
    ["ET", "Extra Trees (Extremely Randomised Trees)", "—"],
    ["CV", "Cross-validation", "—"],
    ["R²", "Coefficient of determination (variance explained)", "—"],
    ["MAE", "Mean absolute error", "(target unit)"],
    ["RMSE", "Root-mean-square error", "(target unit)"],
    ["LOO", "Leave-One-Out cross-validation", "—"],
    ["σ", "Standard deviation (used for the ±1σ uncertainty band)", "—"],
])

# ── Contents note ────────────────────────────────────────────────────────────
h1("Contents")
for c in ["List of Symbols and Abbreviations", "1. Introduction",
          "2. Objectives", "3. Literature Review", "4. Dataset",
          "5. Methodology", "6. The Machine-Learning Model (Detailed)",
          "7. Improving the Model: Before vs After", "8. Results and Discussion",
          "9. Conclusion", "Appendix A: Complete Dataset",
          "Appendix B: Source Code", "References"]:
    bullet(c)

# ── 1. Introduction ──────────────────────────────────────────────────────────
h1("1. Introduction")
h2("1.1 Background")
p("Surface engineering plays a decisive role in the functional performance and "
  "service life of engineering components, because properties such as roughness "
  "and hardness strongly influence wear, fatigue, fluid interaction, and "
  "durability. Laser-based processing has become an important, non-contact way to "
  "tailor these surface and near-surface properties with high precision and "
  "repeatability, and is widely used on metallic materials in the aerospace, "
  "automotive, and biomedical sectors [1], [5].")
h2("1.2 Process Parameters")
p("During laser processing, focused laser energy is delivered to the material "
  "while the beam is scanned across the surface. The outcome is controlled mainly "
  "by four process parameters:")
bullet("Laser power (P) — the energy delivery rate; higher power deposits more "
       "energy into the material.")
bullet("Scan speed (v) — how fast the beam moves; higher speed reduces the energy "
       "received per unit length.")
bullet("Hatch distance (h) — the spacing between adjacent scan tracks, which sets "
       "how much neighbouring tracks overlap.")
bullet("Layer thickness (t) — the depth of material processed per layer.")
p("These four parameters jointly determine the energy delivered per unit volume "
  "of material. Too little energy leaves material poorly consolidated and rough; "
  "too much causes excessive melting and defects [6], [7]. Because the parameters "
  "interact, their combined effect — rather than any single one — governs the "
  "final roughness and hardness.")
h2("1.3 Why Machine Learning")
p("The relationship between these parameters and the resulting properties is "
  "non-linear and involves interacting factors, which makes it difficult to "
  "optimise by trial and error [6], [13]. Earlier studies have used statistical "
  "designs such as the Taguchi method and response-surface methodology to study "
  "these effects [21], [23]; more recently, machine learning has been applied to "
  "predict properties such as roughness, hardness, and density directly from "
  "process parameters [17], [18]. Machine learning learns the mapping directly "
  "from data and predicts properties for new parameter combinations, reducing the "
  "number of physical trials required.")
h2("1.4 Present Work")
p("This project builds two predictive models — one for surface roughness and one "
  "for hardness — from compiled experimental data, improves their accuracy through "
  "feature engineering and model selection, and deploys them in an interactive "
  "tool that reports each prediction with an uncertainty estimate.")

# ── 2. Objectives ────────────────────────────────────────────────────────────
h1("2. Objectives")
bullet("Compile and structure experimental data relating process parameters to "
       "surface roughness and hardness [1]–[16].")
bullet("Train baseline regression models for both properties.")
bullet("Improve accuracy through feature engineering and per-property model "
       "selection.")
bullet("Validate the models honestly using cross-validation (R², MAE, RMSE).")
bullet("Quantify which process parameters most influence each property.")
bullet("Deploy the models in an interactive prediction interface with uncertainty.")

# ── 3. Literature Review ─────────────────────────────────────────────────────
h1("3. Literature Review")
h2("3.1 Process Parameters and Part Properties")
p("A large body of experimental work has characterised how laser process "
  "parameters affect part quality. Relative density and hardness generally "
  "improve with sufficient energy input but degrade when energy is excessive, "
  "while surface roughness depends strongly on layer thickness, hatch distance, "
  "and scan speed [1], [2], [5], [6]. A recurring finding is that no single "
  "parameter explains the response on its own — the parameters interact, and the "
  "combined energy input is often the more meaningful descriptor [11], [19].")
h2("3.2 Statistical and Optimisation Approaches")
p("To reduce experimental effort, researchers have applied statistical designs "
  "such as the Taguchi method and response-surface methodology to identify "
  "influential parameters and near-optimal settings [21], [23]. These methods are "
  "effective for screening and for modest, well-structured experiments, but they "
  "can struggle to capture strongly non-linear interactions across a wide "
  "parameter space.")
h2("3.3 Machine-Learning Approaches")
p("More recently, machine-learning regressors — including tree ensembles such as "
  "Random Forests and related methods — have been used to predict properties like "
  "roughness, hardness, and density from process parameters [17], [18]. Such "
  "models capture non-linear interactions automatically and provide interpretable "
  "feature-importance scores, making them well suited to the modest, multi-source "
  "datasets typical of compiled literature data.")
h2("3.4 Research Gap")
p("While individual studies report parameter–property trends, the data are "
  "scattered across many papers with differing setups, and there is limited "
  "integration of these results into a single, validated predictive tool. This "
  "project addresses that gap by compiling the data, engineering physically "
  "motivated features, and building cross-validated models for both surface "
  "roughness and hardness.")

# ── 4. Dataset ───────────────────────────────────────────────────────────────
h1("4. Dataset")
p("Data were compiled from published experimental studies [1]–[16]. Because the "
  "roughness and hardness measurements come from different experiments, each "
  "property has its own dataset and its own model.")
caption("Table 4.1 — Datasets and target ranges.")
table(["Property", "Samples", "Range", "Cross-validation"],
      [[TARGETS[k]["label"], EV[k]["rows"],
        f"{EV[k]['range'][0]:.1f}–{EV[k]['range'][1]:.1f} {TARGETS[k]['unit']}",
        EV[k]["cv"]] for k in TARGETS])
p("")
caption("Table 4.2 — Input process parameters and their ranges.")
table(["Parameter", "Min", "Max"],
      [[FEATURE_LABELS[f], f"{PARAM_RANGES[f][0]:g}", f"{PARAM_RANGES[f][1]:g}"]
       for f in BASE_FEATURES])

# ── 4. Methodology ───────────────────────────────────────────────────────────
h1("5. Methodology")
h2("5.1 Overview")
p("The end-to-end workflow runs from raw data through feature engineering, model "
  "training, cross-validation, and finally deployment.")
caption("Figure 5.1 — Machine-learning workflow.")
add_picture(RESULTS / "workflow.png", 6.4)

h2("5.2 Feature Engineering")
p("In addition to the four raw parameters, two derived features are computed for "
  "every sample:")
bullet("Energy = P / (v · h · t) — laser energy delivered per unit volume.")
bullet("Power/Speed = P / v — energy delivered per unit track length.")
p("These consolidate the interactions between parameters into single quantities "
  "that the model can use directly, rather than having to reconstruct them from "
  "the raw inputs on a small dataset. As shown in Section 6, adding them improves "
  "accuracy. The implementation is shown below.")
caption("Code — feature engineering (config.py).")
code_block(snippet("config.py", "def add_derived_features", "def build_model"))

h2("5.3 Validation Approach")
p("Each model is assessed with cross-validation, where the data are repeatedly "
  "split so every sample is predicted by a model that did not train on it. For "
  "the small hardness set (<40 samples) Leave-One-Out cross-validation is used; "
  "otherwise 5-fold. Predictions from all held-out folds are pooled and the "
  "metrics — R² (variance explained), MAE (mean absolute error), and RMSE — "
  "computed once over the pooled set. This gives an honest estimate of "
  "generalisation rather than the optimistic numbers a single train/test split "
  "can produce.")

# ── 5. The ML model (detailed) ───────────────────────────────────────────────
h1("6. The Machine-Learning Model (Detailed)")
h2("6.1 Why a Tree-Ensemble Model")
p("The parameter-to-property relationship is non-linear, has interacting factors, "
  "and carries appreciable measurement scatter. Tree-ensemble regressors are "
  "well suited to this setting:")
bullet("They capture non-linear interactions automatically, without manually "
       "specified cross terms [18].")
bullet("As ensembles of many decorrelated trees, they resist over-fitting on "
       "small, noisy datasets [17], [18].")
bullet("They provide interpretable feature-importance scores.")
bullet("The spread across trees gives a built-in uncertainty estimate (±1σ).")
p("Two such algorithms are used. A Random Forest builds each tree on a bootstrap "
  "resample of the data. Extra Trees (Extremely Randomised Trees) additionally "
  "randomises the split thresholds, which further reduces variance and proved "
  "more accurate for the small hardness dataset (Section 6).")

h2("6.2 How the Model Works")
p("Each tree makes its own prediction by following a series of yes/no splits on "
  "the input features. The ensemble's prediction is the average of all the trees' "
  "outputs, which cancels much of the individual trees' error. The standard "
  "deviation across the trees provides the ±1σ uncertainty band reported by the "
  "application.")
caption("Figure 6.1 — A tree ensemble averages many decision trees into one "
        "prediction with an uncertainty band.")
add_picture(RESULTS / "random_forest.png", 6.0)

h2("6.3 Input Features")
caption("Table 6.1 — Model input features.")
table(["#", "Feature", "Type"],
      [[i + 1, FEATURE_LABELS[f], "raw parameter" if f in BASE_FEATURES else "derived"]
       for i, f in enumerate(FEATURES)])

h2("6.4 Model Configuration")
p("Each property uses the algorithm that cross-validated best for it: a tuned "
  "Random Forest for surface roughness and Extra Trees for hardness. Both wrap the "
  "estimator in a pipeline that first standardises the features.")
caption("Table 6.2 — Model configuration per property.")
table(["Property", "Algorithm", "Key settings"], [
    ["Surface Roughness", "Random Forest",
     f"{RF_PARAMS['n_estimators']} trees, min_samples_leaf="
     f"{RF_PARAMS['min_samples_leaf']}, max_features={RF_PARAMS['max_features']}"],
    ["Hardness", "Extra Trees",
     f"{EXTRATREES_PARAMS['n_estimators']} trees, min_samples_leaf="
     f"{EXTRATREES_PARAMS['min_samples_leaf']}"],
])
caption("Code — model definition (config.py).")
code_block(snippet("config.py", "def build_model", max_lines=12))

h2("6.5 Training and Prediction Code")
p("The training routine cross-validates, fits on all data, and saves each model:")
caption("Code — training one property (train.py).")
code_block(snippet("train.py", "def train_one", "def _plot_actual_vs_pred"))
p("At prediction time the app averages the trees and reports the spread as ±1σ:")
caption("Code — prediction with uncertainty (app.py).")
code_block(snippet("app.py", "def predict_one", "# ──", max_lines=10))

h2("6.6 Per-Model Workflows")
p("Although both models share the same pipeline, they differ in the dataset they "
  "learn from, the algorithm chosen, the cross-validation scheme, and the "
  "parameters that drive each property. The two diagrams below trace each model "
  "end-to-end — from input data, through feature engineering and standardisation, "
  "to the specific ensemble algorithm, validation scheme, and final cross-"
  "validated metrics. The side panel lists the top driving features for each.")

h3 = lambda t: doc.add_heading(t, level=3)
h3("6.6.1 Surface Roughness Model (Random Forest)")
p("The surface-roughness model is trained on the 89-sample roughness dataset "
  "using a tuned Random Forest, in which each tree is grown on a bootstrap "
  "resample and the best split is chosen at each node. It is validated with "
  "5-fold cross-validation. Hatch distance dominates the prediction, followed by "
  "the engineered energy and power/speed terms.")
caption("Figure 6.2 — Detailed workflow of the Surface Roughness model.")
add_picture(RESULTS / "workflow_Surface_Roughness_um.png", 5.0)

h3("6.6.2 Hardness Model (Extra Trees)")
p("The hardness model is trained on the smaller 30-sample hardness dataset using "
  "Extra Trees, which randomises the split thresholds to reduce variance — the "
  "key reason it generalises better than a Random Forest on this small set. "
  "Because of the limited sample count it is validated with Leave-One-Out cross-"
  "validation. Layer thickness is the strongest driver, with the power/speed and "
  "energy terms next.")
caption("Figure 6.3 — Detailed workflow of the Hardness model.")
add_picture(RESULTS / "workflow_Hardness_HV.png", 5.0)

# ── 6. Before vs After ───────────────────────────────────────────────────────
h1("7. Improving the Model: Before vs After")
p("The project began with a basic configuration — a single Random Forest using "
  "only the four raw parameters. Two changes were then made and validated:")
bullet("Feature engineering: added the Energy and Power/Speed derived features.")
bullet("Model selection: chose the best algorithm per property — keeping a tuned "
       "Random Forest for roughness, but switching hardness to Extra Trees.")
p("The figure and table below compare the basic and improved configurations under "
  "identical cross-validation.")
caption("Figure 7.1 — Accuracy (R²) and error (MAE): before vs after.")
add_picture(RESULTS / "before_after.png", 6.5)

caption("Table 7.1 — Before vs after (cross-validated).")
rows = []
for k in TARGETS:
    b, a = BA[k]["before"], BA[k]["after"]
    unit = BA[k]["unit"]
    dr2 = a[0] - b[0]
    rows.append([BA[k]["label"],
                 f"{b[0]:.3f}", f"{a[0]:.3f}", f"+{dr2:.3f}",
                 f"{b[1]:.2f} {unit}", f"{a[1]:.2f} {unit}"])
table(["Property", "R² before", "R² after", "Δ R²", "MAE before", "MAE after"], rows)

dh = BA["hardness"]
p(f"The hardness model shows the largest gain: R² rose from {dh['before'][0]:.2f} "
  f"to {dh['after'][0]:.2f} and mean absolute error fell from {dh['before'][1]:.2f} "
  f"to {dh['after'][1]:.2f} HV. Two factors drove this. First, the derived "
  "features gave the model the energy interaction directly. Second, and most "
  "importantly, switching from a Random Forest to Extra Trees helped: with only "
  f"{EV['hardness']['rows']} hardness samples, the extra randomisation of Extra "
  "Trees reduces variance and generalises better than a standard Random Forest. "
  "Surface roughness also improved, more modestly, from feature engineering and "
  "hyperparameter tuning.")

# ── 7. Results and Discussion ────────────────────────────────────────────────
h1("8. Results and Discussion")
h2("8.1 Final Model Performance")
caption("Table 8.1 — Final cross-validated performance.")
table(["Property", "Samples", "CV", "R²", "MAE", "RMSE"],
      [[TARGETS[k]["label"], EV[k]["rows"], EV[k]["cv"], f"{EV[k]['r2']:.3f}",
        f"{EV[k]['mae']:.2f} {TARGETS[k]['unit']}",
        f"{EV[k]['rmse']:.2f} {TARGETS[k]['unit']}"] for k in TARGETS])

h2("8.2 Feature Importance")
caption("Table 8.2 — Feature importance per property.")
trows = []
for f in FEATURES:
    trows.append([FEATURE_LABELS[f]] + [f"{EV[k]['importance'][f]:.3f}" for k in TARGETS])
table(["Feature"] + [TARGETS[k]["label"] for k in TARGETS], trows)
for k, spec in TARGETS.items():
    top = max(EV[k]["importance"], key=EV[k]["importance"].get)
    p(f"For {spec['label'].lower()}, the most influential feature is "
      f"{FEATURE_LABELS[top]} (importance {EV[k]['importance'][top]:.2f}).")

h2("8.3 Property vs Parameters")
for k, spec in TARGETS.items():
    caption(f"Figure 8.{list(TARGETS).index(k)+1} — {spec['label']} vs each "
            f"process parameter.")
    add_picture(RESULTS / f"relational_{spec['column']}.png", 5.6)

h2("8.4 Model Validation Plots")
for k, spec in TARGETS.items():
    caption(f"Figure — {spec['label']}: actual vs predicted (cross-validated).")
    add_picture(MODEL_DIR / f"actual_vs_predicted_{spec['column']}.png", 4.0)

# ── 8. Conclusion ────────────────────────────────────────────────────────────
h1("9. Conclusion")
bullet(f"Two models were built and validated: surface roughness "
       f"(R² = {EV['roughness']['r2']:.2f}) and hardness "
       f"(R² = {EV['hardness']['r2']:.2f}).")
bullet("Feature engineering (energy and power/speed terms) and per-property model "
       "selection measurably improved accuracy.")
bullet(f"The hardness model improved most, from R² = {BA['hardness']['before'][0]:.2f} "
       f"to {BA['hardness']['after'][0]:.2f}, mainly by switching to Extra Trees on "
       "the small dataset.")
bullet("The hardness dataset is small (30 samples); collecting more data is the "
       "clearest route to further improvement.")
bullet("The trained models are deployed in an interactive app that reports "
       "predictions with uncertainty.")

# ── Appendix A: data ─────────────────────────────────────────────────────────
doc.add_page_break()
h1("Appendix A: Complete Dataset")
for key, spec in TARGETS.items():
    raw = pd.read_csv(DATA_DIR / spec["csv"])
    h2(f"A.{list(TARGETS).index(key)+1} {spec['label']} data ({len(raw)} samples)")
    caption(f"All {len(raw)} samples used to train the {spec['label'].lower()} model.")
    headers = ["#", "Power (W)", "Speed (mm/s)", "Hatch (mm)", "Layer (mm)",
               f"{spec['label']} ({spec['unit']})"]
    rows = [[i + 1, f"{r['Laser_Power_W']:g}", f"{r['Scan_Speed_mm_s']:g}",
             f"{r['Hatch_Distance_mm']:g}", f"{r['Layer_Thickness_mm']:g}",
             f"{r[spec['column']]:g}"] for i, (_, r) in enumerate(raw.iterrows())]
    table(headers, rows)
    p("")

# ── Appendix B: code ─────────────────────────────────────────────────────────
doc.add_page_break()
h1("Appendix B: Source Code")
for fname in ["config.py", "train.py", "app.py"]:
    h2(f"B — {fname}")
    code_block((BASE_DIR / fname).read_text(encoding="utf-8"))

# ── References ───────────────────────────────────────────────────────────────
doc.add_page_break()
h1("References")
p("In-text citations use bracketed numbers [n] in IEEE style. References "
  "[1]–[16] are the experimental data sources; [17]–[24] support the modelling "
  "methodology.")
REFERENCES = [
    "A. Maamoun, Y. Xue, M. Elbestawi, and S. Veldhuis, “The Effect of Selective "
    "Laser Melting Process Parameters on the Microstructure and Mechanical "
    "Properties of Al6061 and AlSi10Mg Alloys,” Materials, vol. 12, no. 1, p. 12, "
    "2018, doi: 10.3390/ma12010012.",
    "X. Han, H. Zhu, X. Nie, G. Wang, and X. Zeng, “Investigation on selective "
    "laser melting AlSi10Mg cellular lattice strut: Molten pool morphology, "
    "surface roughness and dimensional accuracy,” Materials, vol. 11, no. 3, 2018, "
    "doi: 10.3390/ma11030392.",
    "P. Wei et al., “The AlSi10Mg samples produced by selective laser melting: "
    "single track, densification, microstructure and mechanical behavior,” Appl "
    "Surf Sci, vol. 408, pp. 38–50, 2017, doi: 10.1016/j.apsusc.2017.02.215.",
    "S. L. Sing, L. P. Lam, D. Q. Zhang, Z. H. Liu, and C. K. Chua, “Interfacial "
    "characterization of SLM parts in multi-material processing: Intermetallic "
    "phase formation between AlSi10Mg and C18400 copper alloy,” Mater Charact, "
    "vol. 107, pp. 220–227, 2015, doi: 10.1016/j.matchar.2015.07.007.",
    "N. Read, W. Wang, K. Essa, and M. M. Attallah, “Selective laser melting of "
    "AlSi10Mg alloy: Process optimisation and mechanical properties development,” "
    "Mater Des, vol. 65, pp. 417–424, 2015, doi: 10.1016/j.matdes.2014.09.044.",
    "K. Kempen, L. Thijs, E. Yasa, M. Badrossamay, W. Verheecke, and J.-P. Kruth, "
    "“Process Optimization and Microstructural Analysis for Selective Laser Melting "
    "of AlSi10Mg,” in 22nd Annual International Solid Freeform Fabrication "
    "Symposium, 2011.",
    "N. T. Aboulkhair, N. M. Everitt, I. Ashcroft, and C. Tuck, “Reducing porosity "
    "in AlSi10Mg parts processed by selective laser melting,” Addit Manuf, vol. 1, "
    "pp. 77–86, 2014, doi: 10.1016/j.addma.2014.08.001.",
    "C. Y. Yap, C. K. Chua, and Z. L. Dong, “An effective analytical model of "
    "selective laser melting,” Virtual Phys Prototyp, vol. 11, no. 1, pp. 21–26, "
    "2016, doi: 10.1080/17452759.2015.1133217.",
    "A. A. Raus, M. S. Wahab, M. Ibrahim, K. Kamarudin, A. Ahmed, and S. Shamsudin, "
    "“Mechanical and physical properties of AlSi10Mg processed through selective "
    "laser melting,” in AIP Conference Proceedings, 2017, doi: 10.1063/1.4981168.",
    "W. H. Kan, Y. Nadot, M. Foley, L. Ridosz, G. Proust, and J. M. Cairney, "
    "“Factors that affect the properties of additively-manufactured AlSi10Mg: "
    "Porosity versus microstructure,” Addit Manuf, vol. 29, 2019, "
    "doi: 10.1016/j.addma.2019.100805.",
    "S. Bai, N. Perevoshchikova, Y. Sha, and X. Wu, “The effects of selective laser "
    "melting process parameters on relative density of the AlSi10Mg parts and "
    "suitable procedures of the archimedes method,” Applied Sciences, vol. 9, "
    "no. 3, 2019, doi: 10.3390/app9030583.",
    "L. Wang, S. Wang, and J. Wu, “Experimental investigation on densification "
    "behavior and surface roughness of AlSi10Mg powders produced by selective "
    "laser melting,” Opt Laser Technol, vol. 96, pp. 88–96, 2017, "
    "doi: 10.1016/j.optlastec.2017.05.006.",
    "O. Poncelet et al., “Critical assessment of the impact of process parameters "
    "on vertical roughness and hardness of thin walls of AlSi10Mg processed by "
    "laser powder bed fusion,” Addit Manuf, vol. 38, p. 101801, 2021, "
    "doi: 10.1016/j.addma.2020.101801.",
    "S. M. Yusuf, M. Hoegden, and N. Gao, “Effect of sample orientation on the "
    "microstructure and microhardness of additively manufactured AlSi10Mg "
    "processed by high-pressure torsion,” Int. J. Adv. Manuf. Technol., vol. 106, "
    "no. 9–10, pp. 4321–4337, 2020, doi: 10.1007/s00170-019-04817-5.",
    "A. Tridello et al., “Effect of microstructure, residual stresses and building "
    "orientation on the fatigue response up to 10^9 cycles of an SLM AlSi10Mg "
    "alloy,” Int J Fatigue, vol. 137, p. 105659, 2020, "
    "doi: 10.1016/j.ijfatigue.2020.105659.",
    "B. J. Mfusi, L. C. Tshabalala, A. P. I. Popoola, and N. R. Mathe, “The effect "
    "of selective laser melting build orientation on the mechanical properties of "
    "AlSi10Mg parts,” IOP Conf Ser Mater Sci Eng, vol. 430, p. 012028, 2018, "
    "doi: 10.1088/1757-899X/430/1/012028.",
    "H. Abdulla, M. Maalouf, I. Barsoum, and H. An, “Truncated Newton Kernel Ridge "
    "Regression for Prediction of Porosity in Additive Manufactured SS316L,” "
    "Applied Sciences, vol. 12, no. 9, 2022, doi: 10.3390/app12094252.",
    "G. O. Barrionuevo, J. A. Ramos-Grez, M. Walczak, and C. A. Betancourt, "
    "“Comparative evaluation of supervised machine learning algorithms in the "
    "prediction of the relative density of 316L stainless steel fabricated by "
    "selective laser melting,” Int. J. Adv. Manuf. Technol., vol. 113, no. 1–2, "
    "pp. 419–433, 2021, doi: 10.1007/s00170-021-06596-4.",
    "C. U. Brown et al., “The effects of laser powder bed fusion process parameters "
    "on material hardness and density for nickel alloy 625,” Gaithersburg, MD, "
    "2018, doi: 10.6028/NIST.AMS.100-19.",
    "D. Wang, J. Lv, X. Wei, D. Lu, and C. Chen, “Study on Surface Roughness "
    "Improvement of Selective Laser Melted Ti6Al4V Alloy,” Crystals, vol. 13, "
    "no. 2, 2023, doi: 10.3390/cryst13020306.",
    "S. Dingal, T. R. Pradhan, J. K. S. Sundar, A. R. Choudhury, and S. K. Roy, "
    "“The application of Taguchi’s method in the experimental investigation of the "
    "laser sintering process,” Int. J. Adv. Manuf. Technol., vol. 38, no. 9–10, "
    "pp. 904–914, 2008, doi: 10.1007/s00170-007-1154-1.",
    "E. Abele, H. A. Stoffregen, M. Kniepkamp, S. Lang, and M. Hampe, “Selective "
    "laser melting for manufacturing of thin-walled porous elements,” J Mater "
    "Process Technol, vol. 215, pp. 114–122, 2015, "
    "doi: 10.1016/j.jmatprotec.2014.07.017.",
    "J. Sun, Y. Yang, and D. Wang, “Parametric optimization of selective laser "
    "melting for forming Ti6Al4V samples by Taguchi method,” Opt Laser Technol, "
    "vol. 49, pp. 118–124, 2013, doi: 10.1016/j.optlastec.2012.12.002.",
    "B. Vandenbroucke and J. Kruth, “Selective laser melting of biocompatible "
    "metals for rapid manufacturing of medical parts,” Rapid Prototyp J, vol. 13, "
    "no. 4, pp. 196–203, 2007, doi: 10.1108/13552540710776142.",
]
for ref in REFERENCES:
    doc.add_paragraph(ref, style="List Number")

# ── Save ─────────────────────────────────────────────────────────────────────
out = "Project_Report.docx"
doc.save(out)
print(f"\nSaved {out}  |  refs={len(REFERENCES)}  |  "
      + "  ".join(f"{k}: {EV[k]['r2']:.3f}" for k in TARGETS))
